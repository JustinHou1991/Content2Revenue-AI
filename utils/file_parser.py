"""
文件解析工具 - 支持 PDF、Word、Excel、CSV 等多种格式
"""

import io
from typing import Dict, Any, List, Optional
import logging

logger = logging.getLogger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """
    解析 PDF 文件，提取文本内容
    
    Args:
        file_bytes: PDF 文件字节流
        
    Returns:
        提取的文本内容
    """
    try:
        import PyPDF2
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text.append(f"--- 第 {page_num + 1} 页 ---\n{page_text}")
            except Exception as e:
                logger.warning(f"提取 PDF 第 {page_num + 1} 页失败: {e}")
                continue
        
        return "\n\n".join(text) if text else ""
    except ImportError:
        logger.error("PyPDF2 未安装，无法解析 PDF")
        raise ImportError("请安装 PyPDF2: pip install PyPDF2")
    except Exception as e:
        logger.error(f"PDF 解析失败: {e}")
        raise


def parse_word(file_bytes: bytes) -> str:
    """
    解析 Word 文件（.docx），提取文本内容
    
    Args:
        file_bytes: Word 文件字节流
        
    Returns:
        提取的文本内容
    """
    try:
        from docx import Document
        doc_file = io.BytesIO(file_bytes)
        doc = Document(doc_file)
        
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n".join(text) if text else ""
    except ImportError:
        logger.error("python-docx 未安装，无法解析 Word")
        raise ImportError("请安装 python-docx: pip install python-docx")
    except Exception as e:
        logger.error(f"Word 解析失败: {e}")
        raise


def parse_excel(file_bytes: bytes, sheet_name: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    解析 Excel 文件，提取数据
    
    Args:
        file_bytes: Excel 文件字节流
        sheet_name: 指定工作表名称，默认第一个
        
    Returns:
        数据列表，每项为字典
    """
    try:
        import pandas as pd
        excel_file = io.BytesIO(file_bytes)
        
        # 读取指定或所有工作表
        if sheet_name:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
        else:
            # 默认读取第一个工作表
            xls = pd.ExcelFile(excel_file)
            first_sheet = xls.sheet_names[0]
            df = pd.read_excel(excel_file, sheet_name=first_sheet)
        
        # 转换为字典列表
        df = df.fillna("")  # 处理空值
        records = df.to_dict("records")
        return records
    except ImportError:
        logger.error("pandas/openpyxl 未安装，无法解析 Excel")
        raise ImportError("请安装: pip install pandas openpyxl")
    except Exception as e:
        logger.error(f"Excel 解析失败: {e}")
        raise


def parse_csv(file_bytes: bytes) -> List[Dict[str, Any]]:
    """
    解析 CSV 文件，提取数据
    
    Args:
        file_bytes: CSV 文件字节流
        
    Returns:
        数据列表，每项为字典
    """
    try:
        import pandas as pd
        csv_file = io.BytesIO(file_bytes)
        df = pd.read_csv(csv_file)
        df = df.fillna("")
        records = df.to_dict("records")
        return records
    except Exception as e:
        logger.error(f"CSV 解析失败: {e}")
        raise


def parse_text(file_bytes: bytes, encoding: str = "utf-8") -> str:
    """
    解析纯文本文件
    
    Args:
        file_bytes: 文本文件字节流
        encoding: 编码格式
        
    Returns:
        文本内容
    """
    try:
        return file_bytes.decode(encoding)
    except UnicodeDecodeError:
        # 尝试其他编码
        for enc in ["gbk", "gb2312", "latin-1"]:
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        raise ValueError("无法识别文件编码")


def parse_file(uploaded_file) -> Dict[str, Any]:
    """
    通用文件解析入口，根据文件扩展名自动选择解析器
    
    Args:
        uploaded_file: Streamlit 的 UploadedFile 对象
        
    Returns:
        {
            "type": "text" | "table",
            "content": str | List[Dict],
            "filename": str,
            "file_type": str
        }
    """
    if uploaded_file is None:
        raise ValueError("未提供文件")
    
    filename = uploaded_file.name
    file_type = filename.lower().split(".")[-1] if "." in filename else ""
    file_bytes = uploaded_file.getvalue()
    
    result = {
        "filename": filename,
        "file_type": file_type,
        "content": None,
        "type": None
    }
    
    # 根据文件类型选择解析器
    if file_type == "pdf":
        result["content"] = parse_pdf(file_bytes)
        result["type"] = "text"
        
    elif file_type in ["docx", "doc"]:
        result["content"] = parse_word(file_bytes)
        result["type"] = "text"
        
    elif file_type in ["xlsx", "xls"]:
        result["content"] = parse_excel(file_bytes)
        result["type"] = "table"
        
    elif file_type == "csv":
        result["content"] = parse_csv(file_bytes)
        result["type"] = "table"
        
    elif file_type in ["txt", "md", "json"]:
        result["content"] = parse_text(file_bytes)
        result["type"] = "text"
        
    else:
        raise ValueError(f"不支持的文件格式: {file_type}")
    
    return result


def extract_text_for_analysis(parse_result: Dict[str, Any], max_length: int = 5000) -> str:
    """
    从解析结果中提取文本用于 AI 分析
    
    Args:
        parse_result: parse_file 的返回结果
        max_length: 最大文本长度
        
    Returns:
        用于分析的文本
    """
    content = parse_result.get("content", "")
    
    if parse_result["type"] == "text":
        text = content
    elif parse_result["type"] == "table":
        # 表格数据转换为文本描述
        if isinstance(content, list) and len(content) > 0:
            # 取前几行作为示例
            sample_rows = content[:5]
            text_parts = []
            for i, row in enumerate(sample_rows, 1):
                row_desc = ", ".join([f"{k}: {v}" for k, v in row.items() if v])
                text_parts.append(f"记录 {i}: {row_desc}")
            
            if len(content) > 5:
                text_parts.append(f"... 共 {len(content)} 条记录")
            
            text = "\n".join(text_parts)
        else:
            text = str(content)
    else:
        text = str(content)
    
    # 截断超长文本
    if len(text) > max_length:
        text = text[:max_length] + f"\n\n[内容已截断，原长度 {len(text)} 字符]"
    
    return text
