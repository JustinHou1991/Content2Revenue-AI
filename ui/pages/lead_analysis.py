"""
线索分析页面 - 销售线索画像构建
使用新设计系统组件 (design_system.py + styles.py)
继承自 AnalysisPage 基类
"""

import streamlit as st
import threading
import logging
import pandas as pd

logger = logging.getLogger(__name__)

from ui.base_page import AnalysisPage
from ui.components.forms import render_lead_form, render_action_buttons
from ui.components.data_display import (
    render_lead_profile_details,
    render_list_section,
)
from ui.components.design_system import (
    page_header,
    metric_card,
    divider,
    callout,
    empty_state,
)
from ui.styles import COLORS

MAX_CSV_SIZE = 10 * 1024 * 1024  # 10MB


class LeadAnalysisPage(AnalysisPage):
    """线索分析页面类"""

    def __init__(self):
        super().__init__(
            title="线索智能分析",
            icon="&#128100;",
            description="分析销售线索，构建客户画像：行业、痛点、购买阶段、意向度等",
            page_prefix="lead"
        )

    def _render_single_input(self):
        """渲染单个线索录入界面"""
        st.subheader("录入线索信息")

        lead_data = render_lead_form()

        analyze_btn, _ = render_action_buttons("开始分析")

        if analyze_btn:
            if not lead_data.get("company") and not lead_data.get("conversation"):
                callout("请至少填写公司名称或对话记录", type="warning")
                return

            with st.spinner("AI正在分析线索..."):
                try:
                    result = self._get_orchestrator().analyze_lead(lead_data)
                    self._display_result(result)
                except Exception as e:
                    callout(f"分析失败: {str(e)}", type="error")
                    st.info("请检查API Key是否有效，或稍后重试。")

    def _render_batch_input(self):
        st.subheader("批量导入线索")

        uploaded_file = st.file_uploader(
            "上传文件（支持 CSV、Excel、Word、PDF）",
            type=["csv", "xlsx", "xls", "docx", "doc", "pdf"],
            key="lead_batch_file"
        )

        if "lead_field_mapping" not in st.session_state:
            st.session_state.lead_field_mapping = None
        if "lead_df" not in st.session_state:
            st.session_state.lead_df = None

        if uploaded_file is not None:
            if uploaded_file.size > MAX_CSV_SIZE:
                callout(f"文件大小超过限制（最大 {MAX_CSV_SIZE // (1024*1024)}MB）", type="error")
                return

            from utils.field_mapping import (
                detect_columns,
                show_mapping_preview,
                validate_mapping_for_analysis,
            )

            if st.session_state.lead_df is None:
                import io
                import csv
                file_type = uploaded_file.name.lower().split(".")[-1]

                if file_type in ["xlsx", "xls"]:
                    st.session_state.lead_df = pd.read_excel(uploaded_file)
                elif file_type == "pdf":
                    st.session_state.lead_df = self._parse_pdf(uploaded_file)
                elif file_type in ["docx", "doc"]:
                    st.session_state.lead_df = self._parse_word(uploaded_file)
                else:
                    file_bytes = uploaded_file.getvalue()
                    csv_file = io.BytesIO(file_bytes)

                    success = False
                    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
                        if success:
                            break
                        for sep in [",", ";", "\t", "|"]:
                            try:
                                csv_file.seek(0)
                                df = pd.read_csv(csv_file, encoding=encoding, sep=sep, engine="python", on_bad_lines="skip")
                                if len(df.columns) > 1:
                                    st.session_state.lead_df = df
                                    success = True
                                    break
                            except Exception:
                                continue

                    if not success:
                        st.error("无法解析CSV文件，请检查文件格式或尝试使用Excel格式(.xlsx)")
                        return

            df = st.session_state.lead_df

            if df is None or len(df) == 0:
                callout("文件解析失败或无有效内容", type="error")
                return

            st.markdown("---")
            st.subheader("📋 选择线索描述列")
            st.caption("系统需要知道哪一列包含线索的描述信息（如需求描述、对话记录等）")

            from utils.field_mapping import REVERSE_MAPPING
            auto_col = None
            for col in df.columns:
                col_lower = str(col).lower().strip()
                if col_lower in REVERSE_MAPPING and REVERSE_MAPPING[col_lower] == "需求描述":
                    auto_col = col
                    break
                for keyword in ["需求", "描述", "对话", "记录", "内容", "备注", "说明"]:
                    if keyword in col_lower:
                        auto_col = col
                        break
                if auto_col:
                    break

            if auto_col is None:
                for col in df.columns:
                    if df[col].dtype == "object" and df[col].notna().any():
                        sample = str(df[col].dropna().iloc[0]) if len(df[col].dropna()) > 0 else ""
                        if len(sample) > 20:
                            auto_col = col
                            break

            col_options = ["-- 请选择 --"] + list(df.columns)
            default_idx = col_options.index(auto_col) if auto_col and auto_col in col_options else 0

            selected_col = st.selectbox(
                "📌 主分析列（必需）",
                col_options,
                index=default_idx,
                help="选择包含客户主要需求/留言内容的列，系统将以此列内容作为分析基础",
            )

            other_cols = [c for c in df.columns if c != selected_col and c != "-- 请选择 --"]
            reference_cols = st.multiselect(
                "📋 参考列（可选）",
                options=other_cols,
                default=[],
                help="选择额外的参考列（如公司名称、联系人、电话等），系统将结合多列内容进行综合分析。不选择则仅使用主分析列。",
            )

            if selected_col != "-- 请选择 --":
                all_selected_cols = [selected_col] + reference_cols
                st.session_state.lead_field_mapping = {"需求描述": selected_col, "参考列": reference_cols}
                
                cols_info = f"主列「{selected_col}」"
                if reference_cols:
                    cols_info += f" + {len(reference_cols)} 个参考列"
                st.success(f"✅ 已设置分析列（共 {len(df)} 条记录）: {cols_info}")

                with st.expander("👁️ 预览数据内容"):
                    st.write("**主分析列预览：**")
                    preview_data = df[selected_col].head(5).tolist()
                    for i, val in enumerate(preview_data, 1):
                        val_str = str(val)[:100] if val else "(空)"
                        st.write(f"{i}. {val_str}")
                    
                    if reference_cols:
                        st.write("**参考列预览：**")
                        for ref_col in reference_cols[:3]:
                            ref_vals = df[ref_col].head(3).tolist()
                            st.write(f"- {ref_col}: {ref_vals}")
                    
                    system_keywords = ["在抖音", "发起了", "输入了手机号"]
                    has_system_msg = any(any(kw in str(v) for kw in system_keywords) for v in preview_data if v)
                    if has_system_msg:
                        st.warning("⚠️ 检测到抖音系统消息，这些不是客户需求。如果主分析列选择错误，请重新选择包含客户留言的列。")

                batch_btn = st.button(
                    f"🚀 开始批量分析（{len(df)} 条线索）",
                    type="primary",
                    use_container_width=True,
                )

                if batch_btn:
                    self._handle_batch_analysis()
            else:
                st.warning("请选择包含线索描述的列")
        else:
            st.session_state.lead_df = None
            st.session_state.lead_field_mapping = None

    def _handle_batch_analysis(self):
        import time
        import pandas as pd

        if st.session_state.lead_field_mapping is None:
            callout("请先完成字段映射", type="error")
            return

        mapping = st.session_state.lead_field_mapping
        df = st.session_state.lead_df.copy()

        if "需求描述" not in mapping:
            callout("缺少必需的'需求描述'字段映射", type="error")
            return

        desc_col = mapping.get("需求描述")
        reference_cols = mapping.get("参考列", [])

        if not desc_col or desc_col not in df.columns:
            callout(f"主分析列不存在，请重新选择", type="error")
            return

        st.subheader("数据清洗")

        df_before_total = len(df)

        all_cols_to_clean = [desc_col] + reference_cols
        for col in all_cols_to_clean:
            if col in df.columns:
                df[col] = df[col].astype(str).str.strip()
                df[col] = df[col].replace(
                    ['nan', 'None', 'null', 'NULL', 'Nan', 'NAN', '-', '--', '无', '', 'nat'],
                    ''
                )

        empty_mask = (
            (df[desc_col] == '') |
            (df[desc_col].str.lower().isin(['nan', 'none', 'null', '无', '-', '--', 'nat'])) |
            (df[desc_col].isna())
        )
        empty_count = empty_mask.sum()

        df = df[~empty_mask].reset_index(drop=True)

        cleaned_rows = df_before_total - len(df)

        if cleaned_rows > 0:
            st.success(f"✅ 数据清洗完成：共 {df_before_total} 行，删除 {cleaned_rows} 行无效数据，剩余 {len(df)} 行")
        else:
            st.info(f"数据清洗完成：共 {len(df)} 行数据，无需清洗")

        with st.expander("查看清洗后的数据"):
            preview_cols = [desc_col] + reference_cols[:3]
            preview_cols = [c for c in preview_cols if c in df.columns]
            st.dataframe(df[preview_cols].head(10))

        leads = []
        seen_texts = set()
        skip_reasons = {"empty": 0, "nan": 0, "duplicate": 0, "system_msg": 0}
        
        system_keywords = [
            "在抖音私信输入了手机号",
            "在抖音企业主页发起了通话",
            "在抖音私信发起了通话",
            "在抖音主页发起了咨询",
            "通过抖音广告进入",
        ]
        
        for idx, row in df.iterrows():
            conversation_parts = []
            
            main_content = str(row.get(desc_col, "")).strip()
            if main_content and main_content.lower() not in ["nan", "none", "null", ""]:
                conversation_parts.append(f"[{desc_col}] {main_content}")
            
            for ref_col in reference_cols:
                if ref_col in df.columns:
                    ref_content = str(row.get(ref_col, "")).strip()
                    if ref_content and ref_content.lower() not in ["nan", "none", "null", ""]:
                        conversation_parts.append(f"[{ref_col}] {ref_content}")
            
            conversation = " | ".join(conversation_parts)
            
            if not conversation or conversation.strip() == "":
                skip_reasons["empty"] += 1
                continue
            if conversation.lower() in ["nan", "none", "", "null"]:
                skip_reasons["nan"] += 1
                continue
            
            if main_content:
                is_system_msg = any(keyword in main_content for keyword in system_keywords)
                if is_system_msg:
                    skip_reasons["system_msg"] += 1
                    continue

            conv_key = main_content.strip()[:50] if main_content else conversation.strip()[:50]
            if conv_key in seen_texts:
                skip_reasons["duplicate"] += 1
                continue
            seen_texts.add(conv_key)

            lead_data = {"conversation": conversation}
            for col_name in df.columns:
                if col_name == desc_col:
                    continue
                val = row.get(col_name, "")
                if pd.notna(val) and str(val).strip() and str(val).strip().lower() not in ["nan", "none", ""]:
                    lead_data[col_name] = str(val).strip()

            leads.append({
                "lead_data": lead_data,
                "lead_id": str(idx),
            })
        
        total_skipped = sum(skip_reasons.values())
        if total_skipped > 0:
            with st.expander("🔍 数据提取详情"):
                st.write(f"跳过原因统计：")
                st.write(f"- 空内容: {skip_reasons['empty']} 行")
                st.write(f"- NaN/None/Null: {skip_reasons['nan']} 行")
                st.write(f"- 重复内容: {skip_reasons['duplicate']} 行")
                st.write(f"- 抖音系统消息: {skip_reasons['system_msg']} 行")
                if skip_reasons['system_msg'] > 0:
                    st.info("💡 提示：大量数据被识别为抖音系统消息（如'在抖音私信输入了手机号'）。请检查是否选择了包含客户实际需求的列，而不是'最新互动记录'系统列。")
                if leads:
                    st.write(f"---")
                    st.write(f"✅ 前5条被提取的线索:")
                    for lead in leads[:5]:
                        st.write(f"- {lead['lead_data'].get('conversation', '')[:60]}...")
                else:
                    st.warning("⚠️ 没有提取到任何有效线索。请确认上传的文件中包含客户的实际需求描述（如留言、咨询内容），而不是系统操作记录。")

        total_rows = len(df)
        extracted_count = len(leads)
        skipped_count = total_rows - extracted_count

        st.info(f"数据提取统计: 共 {total_rows} 行, 有效 {extracted_count} 行, 跳过 {skipped_count} 行")

        if not leads:
            callout("未找到有效的线索数据", type="error")
            return

        total = len(leads)

        st.warning("⚠️ **分析中，请勿切换页面**，完成后将自动显示结果。")
        st.info(f"共 {total} 条线索，并发处理中...")

        progress_bar = st.progress(0, text=f"准备分析 {total} 条线索...")
        status_text = st.empty()

        orchestrator = self._get_orchestrator()

        results = [None] * total
        completed = 0

        max_workers = min(3, total)

        def analyze_one(index: int, lead: dict):
            try:
                result = orchestrator.lead_analyzer.analyze(
                    lead_data=lead.get("lead_data", {}),
                    lead_id=lead.get("lead_id"),
                )
                orchestrator.db.save_lead_analysis(result)
                return index, {"success": True, "data": result}
            except Exception as e:
                logger.error(f"线索分析失败 (item {index+1}/{total}): {e}")
                return index, {"success": False, "error": str(e)}

        from concurrent.futures import ThreadPoolExecutor, as_completed
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(analyze_one, i, leads[i]): i
                for i in range(total)
            }
            for future in as_completed(futures):
                idx, result = future.result()
                results[idx] = result
                completed += 1
                pct = int(completed / total * 100)
                progress_bar.progress(
                    pct / 100,
                    text=f"分析中 {completed}/{total} ({pct}%)"
                )
                status_text.info(f"⏳ 已完成 {completed}/{total} 条")

        state = {
            "results": results,
            "total": total,
        }

        self._show_lead_batch_results(state)

    def _show_lead_batch_results(self, state: dict):
        results = state.get("results", [])
        total = state.get("total", len(results))
        success_count = sum(1 for r in results if r.get("success"))
        fail_count = sum(1 for r in results if not r.get("success"))

        msg = f"批量分析完成！成功 {success_count}/{total} 条"
        if fail_count > 0:
            msg += f"（{fail_count} 条失败）"
        callout(msg, type="success", icon="&#10003;")

        divider()
        st.subheader("分析结果")

        for i, r in enumerate(results):
            if r.get("success"):
                data = r.get("data", {})
                profile = data.get("profile", data)
                raw = data.get("raw_data", data.get("raw_data_json", {}))
                lid = data.get("lead_id", "")[:8]
                company = raw.get("company", raw.get("公司名称", ""))
                name = raw.get("name", raw.get("联系人", ""))
                score = profile.get("lead_score", 0)
                grade = profile.get("lead_grade", "N/A")

                label = f"{company}" if company else f"线索"
                if name:
                    label += f" · {name}"

                with st.expander(
                    f"#{i+1} [{lid}] {'🟢' if grade in ['A','B+'] else '🟡' if grade == 'B' else '🔴'} {grade}级 | {score}分 | {label}"
                ):
                    self._display_profile_simple(profile)
            else:
                with st.expander(f"线索 #{i+1} - 分析失败"):
                    st.error(r.get("error", "未知错误"))

        st.session_state.lead_df = None
        st.session_state.lead_field_mapping = None
        if "lead_batch_state" in st.session_state:
            del st.session_state.lead_batch_state

    def _display_result(self, result: dict):
        """展示单个分析结果（优化版：层次分明）"""
        profile = result["profile"]
        lead_id = result["lead_id"]
        raw = result.get("raw_data", {})

        # === 顶部：线索来源标识 ===
        company = raw.get("company", raw.get("公司名称", ""))
        name = raw.get("name", raw.get("联系人", ""))
        source_label = f"{company}"
        if name:
            source_label += f" · {name}"
        if not source_label:
            source_label = f"线索 {lead_id[:8]}"

        st.markdown(f"### 📋 {source_label}")
        st.caption(f"ID: {lead_id[:8]}... | 分析时间: {result.get('created_at', '未知')[:16]}")

        # === 核心指标卡片（3个最重要的） ===
        col1, col2, col3 = st.columns(3)
        with col1:
            score = profile.get('lead_score', 0)
            score_color = "#10B981" if score >= 70 else "#F59E0B" if score >= 50 else "#EF4444"
            metric_card(
                title="线索评分",
                value=f"{score}/100",
                subtitle="综合质量评分",
                icon="&#128200;",
                border_color=score_color,
            )
        with col2:
            metric_card(
                title="线索等级",
                value=profile.get('lead_grade', 'N/A'),
                subtitle="A(85+) B+(70+) B(55+) C(40+) D(<40)",
                icon="&#127942;",
                border_color="#10B981",
            )
        with col3:
            intent = profile.get('intent_level', 0)
            intent_color = "#10B981" if intent >= 7 else "#F59E0B" if intent >= 4 else "#EF4444"
            metric_card(
                title="购买意向",
                value=f"{intent}/10",
                subtitle="高意向" if intent >= 7 else "中意向" if intent >= 4 else "低意向",
                icon="&#128065;",
                border_color=intent_color,
            )

        divider()

        # === 关键信息（2列布局） ===
        col_left, col_right = st.columns(2)

        with col_left:
            st.markdown("#### 🏢 基础画像")
            info_items = [
                ("行业", profile.get('industry', '未知')),
                ("公司阶段", profile.get('company_stage', '未知')),
                ("决策角色", profile.get('role', '未知')),
                ("购买阶段", profile.get('buying_stage', '未知')),
                ("紧迫程度", profile.get('urgency', '未知')),
                ("预算准备", profile.get('budget_readiness', '未知')),
            ]
            for label, value in info_items:
                st.write(f"**{label}**: {value}")

        with col_right:
            st.markdown("#### 💔 核心痛点")
            pain_points = profile.get("pain_points", [])
            if pain_points:
                for pain in pain_points:
                    st.write(f"- {pain}")
            else:
                st.write("暂无")

            st.markdown("#### 📡 意向信号")
            signals = profile.get("intent_signals", [])
            if signals:
                for signal in signals:
                    st.write(f"- {signal}")
            else:
                st.write("暂无")

        divider()

        # === 策略建议（折叠） ===
        with st.expander("💡 互动策略建议", expanded=False):
            st.info(profile.get("engagement_strategy", "暂无建议"))
            col_a, col_b = st.columns(2)
            with col_a:
                st.write(f"**推荐内容类型**: {profile.get('recommended_content_type', '未知')}")
            with col_b:
                st.write(f"**推荐CTA类型**: {profile.get('recommended_cta', '未知')}")

        with st.expander("⚖️ 决策标准 & 异议风险", expanded=False):
            col_a, col_b = st.columns(2)
            with col_a:
                render_list_section("决策标准", profile.get("decision_criteria", []))
            with col_b:
                render_list_section("异议风险", profile.get("objection_risks", []))

    def _display_profile_simple(self, profile: dict):
        """展示线索画像（简化版，用于历史记录）"""
        col_left, col_right = st.columns(2)

        with col_left:
            st.markdown("**📋 线索评估**")
            # 联系方式
            has_contact = profile.get('has_contact_info', '否')
            contact_type = profile.get('contact_type', '无')
            st.write(f"**联系方式**: {contact_type} ({has_contact})")
            # 线索质量与优先级
            quality = profile.get('lead_quality', '低')
            priority = profile.get('follow_up_priority', '低')
            is_valid = profile.get('is_valid_lead', '否')
            st.write(f"**质量/优先级**: {quality} / {priority} | 有效: {is_valid}")
            # 需求
            req = profile.get('requirement', '未详细说明')
            st.write(f"**需求**: {req}")

        with col_right:
            st.markdown("**🎯 意向评估**")
            # 意向度
            intent = profile.get('intent_level', 5)
            sat = profile.get('satisfaction_level', 5)
            st.write(f"**意向度/满意度**: {intent}/10 / {sat}/10")
            # 核心痛点
            pains = profile.get("pain_points", [])
            if pains:
                for pain in pains[:2]:  # 最多显示2个
                    st.write(f"- {pain}")
            else:
                st.write("暂无痛点信息")

    def _parse_pdf(self, uploaded_file) -> "pd.DataFrame":
        """解析 PDF 文件，提取文本内容"""
        import pandas as pd

        parse_mode = st.radio(
            "PDF 解析模式",
            options=["按页提取（推荐）", "表格提取"],
            horizontal=True,
            help="按页提取：每页作为一条线索；表格提取：尝试识别PDF中的表格结构",
        )

        try:
            if parse_mode == "表格提取":
                return self._parse_pdf_as_table(uploaded_file)
            else:
                return self._parse_pdf_as_text(uploaded_file)
        except Exception as e:
            st.error(f"PDF 解析失败: {str(e)}")
            return pd.DataFrame()

    def _parse_pdf_as_text(self, uploaded_file) -> "pd.DataFrame":
        """按页提取 PDF 文本"""
        import pandas as pd

        try:
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                texts = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append({"需求描述": text.strip(), "页码": i + 1})
        except ImportError:
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(uploaded_file)
            texts = []
            for i, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    texts.append({"需求描述": text.strip(), "页码": i + 1})

        if not texts:
            st.warning("PDF 文件未提取到有效文本内容")
            return pd.DataFrame()

        df = pd.DataFrame(texts)
        st.success(f"✅ 成功从 PDF 提取 {len(df)} 页内容")
        return df

    def _parse_pdf_as_table(self, uploaded_file) -> "pd.DataFrame":
        """尝试从 PDF 提取表格"""
        import pandas as pd

        try:
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                all_tables = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        cleaned = []
                        for row in table:
                            cleaned_row = [(cell or "").strip() for cell in row]
                            if any(cell for cell in cleaned_row):
                                cleaned.append(cleaned_row)
                        if cleaned:
                            all_tables.extend(cleaned)
        except ImportError:
            st.error("表格模式需要 pdfplumber 库")
            return pd.DataFrame()

        if not all_tables or len(all_tables) <= 1:
            st.warning("未检测到有效表格，请尝试按页提取模式")
            return pd.DataFrame()

        # 处理不规则表格（列数不一致）
        max_cols = max(len(row) for row in all_tables)
        normalized = []
        for row in all_tables:
            if len(row) < max_cols:
                row = row + [""] * (max_cols - len(row))
            normalized.append(row[:max_cols])

        headers = normalized[0]
        df = pd.DataFrame(normalized[1:], columns=headers)
        df = df.dropna(how="all")
        df = df[~df.apply(lambda row: all(str(v).strip() in ["", "nan", "None"] for v in row), axis=1)]

        if len(df) == 0:
            st.warning("表格数据为空")
            return pd.DataFrame()

        st.success(f"✅ 成功从 PDF 表格提取 {len(df)} 条记录")
        return df

    def _parse_word(self, uploaded_file) -> "pd.DataFrame":
        """解析 Word 文件，提取文本内容"""
        import pandas as pd
        from docx import Document

        try:
            doc = Document(uploaded_file)
            has_tables = len(doc.tables) > 0

            if has_tables:
                parse_mode = st.radio(
                    "Word 解析模式",
                    options=["表格提取（推荐）", "段落提取"],
                    horizontal=True,
                    help="表格提取：提取文档中的表格数据；段落提取：按段落提取文本",
                )
            else:
                parse_mode = "段落提取"

            if parse_mode == "表格提取":
                return self._parse_word_as_table(doc)
            else:
                return self._parse_word_as_paragraphs(doc)

        except Exception as e:
            st.error(f"Word 解析失败: {str(e)}")
            return pd.DataFrame()

    def _parse_word_as_table(self, doc) -> "pd.DataFrame":
        """从 Word 文档提取表格数据"""
        import pandas as pd

        tables_data = []
        for table in doc.tables:
            for row in table.rows:
                seen_texts = set()
                row_data = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in seen_texts:
                        seen_texts.add(text)
                        row_data.append(text)
                if any(row_data):
                    tables_data.append(row_data)

        if not tables_data or len(tables_data) <= 1:
            st.warning("未检测到有效表格，请尝试段落提取模式")
            return pd.DataFrame()

        # 处理不规则表格（列数不一致）
        max_cols = max(len(row) for row in tables_data)
        normalized = []
        for row in tables_data:
            if len(row) < max_cols:
                row = row + [""] * (max_cols - len(row))
            normalized.append(row[:max_cols])

        headers = normalized[0]
        df = pd.DataFrame(normalized[1:], columns=headers)
        df = df.dropna(how="all")
        df = df[~df.apply(lambda row: all(str(v).strip() in ["", "nan", "None"] for v in row), axis=1)]

        if len(df) == 0:
            st.warning("表格数据为空")
            return pd.DataFrame()

        st.success(f"✅ 成功从 Word 表格提取 {len(df)} 条记录")
        return df

    def _parse_word_as_paragraphs(self, doc) -> "pd.DataFrame":
        """从 Word 文档提取段落文本（智能合并连续短段落）"""
        import pandas as pd

        paragraphs = []
        buffer = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if buffer and len(buffer) > 20:
                    paragraphs.append(buffer.strip())
                buffer = ""
                continue

            if len(text) < 30 and len(buffer) < 200:
                buffer += text
            else:
                if buffer and len(buffer) > 20:
                    paragraphs.append(buffer.strip())
                buffer = text

        if buffer and len(buffer) > 20:
            paragraphs.append(buffer.strip())

        if not paragraphs:
            st.warning("Word 文档未提取到有效文本内容")
            return pd.DataFrame()

        df = pd.DataFrame({"需求描述": paragraphs})
        st.success(f"✅ 成功从 Word 提取 {len(df)} 个段落")
        return df

    def _render_history(self):
        """展示历史记录"""
        st.subheader("历史分析记录")
        try:
            # 分页参数
            page_size = 10
            page = self._get_current_page("lead_history")

            # 获取总数
            total_count = self._get_orchestrator().db.get_lead_analyses_count()

            if total_count == 0:
                empty_state(
                    title="暂无历史记录",
                    description="去上方录入线索开始分析吧！",
                    icon="&#128101;",
                )
                return

            # 获取当前页数据
            offset = page * page_size
            records = self._get_orchestrator().db.get_all_lead_analyses(
                limit=page_size, offset=offset
            )

            if not records:
                empty_state(
                    title="暂无历史记录",
                    description="没有找到匹配的记录。",
                    icon="&#128101;",
                )
                return

            for record in records:
                profile = record.get("profile_json", {})
                score = profile.get("lead_score", 0)
                grade = profile.get("lead_grade", "N/A")
                industry = profile.get("industry", "未知")
                raw = record.get("raw_data_json", {})
                company = raw.get("company", raw.get("公司名称", ""))
                name = raw.get("name", raw.get("联系人", ""))
                intent = profile.get("intent_level", 0)

                # 标题：等级 + 评分 + 公司 + 行业
                title_parts = [f"{grade}级", f"{score}分"]
                if company:
                    title_parts.append(company)
                if industry and industry != "未知":
                    title_parts.append(industry)
                title = " | ".join(title_parts)
                rid = record.get("id", "")[:8]

                # 颜色标识
                grade_icon = "🟢" if grade in ["A", "B+"] else "🟡" if grade == "B" else "🔴"

                with st.expander(f"[{rid}] {grade_icon} {title}"):
                    # 简要信息
                    if name:
                        st.write(f"**联系人**: {name}")
                    self._display_profile_simple(profile)

            # 分页控制
            self._render_pagination(total_count, page_size, "lead_history")

        except Exception as e:
            callout(f"加载历史记录失败: {str(e)}", type="error")


def render_lead_analysis():
    """页面入口函数"""
    page = LeadAnalysisPage()
    page.render()
