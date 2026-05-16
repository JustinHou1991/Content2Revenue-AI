"""
内容分析页面 - 抖音脚本特征提取
使用新设计系统组件 (design_system.py + styles.py)
继承自 AnalysisPage 基类
"""

import streamlit as st
import threading
import logging
import pandas as pd

logger = logging.getLogger(__name__)

from ui.base_page import AnalysisPage
from ui.components.forms import render_text_area, render_action_buttons
from ui.components.data_display import (
    render_content_analysis_details,
    render_list_section,
    render_tags,
)
from ui.components.design_system import (
    page_header,
    metric_card,
    status_badge,
    divider,
    callout,
    empty_state,
)
from ui.styles import COLORS

MAX_CSV_SIZE = 10 * 1024 * 1024  # 10MB


class ContentAnalysisPage(AnalysisPage):
    """内容分析页面类"""

    def __init__(self):
        super().__init__(
            title="内容智能分析",
            icon="📱",
            description="分析抖音脚本，提取Hook类型、情感基调、叙事结构、CTA等特征",
            page_prefix="content"
        )

    def _render_single_input(self):
        """渲染单个脚本输入界面（支持文件上传）"""
        st.subheader("输入抖音脚本")

        # 文件上传
        uploaded_file = st.file_uploader(
            "上传文件（支持 PDF、Word、TXT）",
            type=["pdf", "docx", "doc", "txt", "md"],
            key="content_single_file"
        )

        # 如果有文件上传，解析内容
        file_text = ""
        if uploaded_file is not None:
            try:
                from utils.file_parser import parse_file, extract_text_for_analysis
                parse_result = parse_file(uploaded_file)
                file_text = extract_text_for_analysis(parse_result, max_length=3000)
                st.success(f"✅ 文件「{uploaded_file.name}」解析成功")
            except Exception as e:
                st.error(f"文件解析失败: {str(e)}")
                st.info("支持的格式：PDF、Word(.docx)、TXT、Markdown")

        # 文本输入框（文件内容自动填充）
        default_text = file_text if file_text else ""
        script_text = render_text_area(
            "粘贴脚本内容或直接编辑上传的文件内容",
            value=default_text,
            placeholder="例：你是不是还在用传统方式获客？每天花500块投流，一个询盘都没有？...",
            height=200
        )

        analyze_btn, _ = render_action_buttons("开始分析")

        if analyze_btn:
            if not script_text:
                callout("请先输入脚本内容或上传文件", type="warning")
                return

            with st.spinner("AI正在分析脚本..."):
                try:
                    result = self._get_orchestrator().analyze_content(script_text)
                    self._display_result(result)
                except Exception as e:
                    callout(f"分析失败: {str(e)}", type="error")
                    st.info("请检查API Key是否有效，或稍后重试。")

    def _render_batch_input(self):
        st.subheader("批量导入脚本")

        uploaded_file = st.file_uploader(
            "上传文件（支持 CSV、Excel、Word、PDF）",
            type=["csv", "xlsx", "xls", "docx", "doc", "pdf"],
            key="content_batch_file"
        )

        if "content_field_mapping" not in st.session_state:
            st.session_state.content_field_mapping = None
        if "content_df" not in st.session_state:
            st.session_state.content_df = None

        if uploaded_file is not None:
            if uploaded_file.size > MAX_CSV_SIZE:
                callout(f"文件大小超过限制（最大 {MAX_CSV_SIZE // (1024*1024)}MB）", type="error")
                return

            import pandas as pd
            from utils.field_mapping import (
                detect_columns,
                show_mapping_preview,
                validate_mapping_for_analysis,
            )

            if st.session_state.content_df is None:
                import io
                file_type = uploaded_file.name.lower().split(".")[-1]

                if file_type in ["xlsx", "xls"]:
                    st.session_state.content_df = pd.read_excel(uploaded_file)
                elif file_type == "pdf":
                    st.session_state.content_df = self._parse_pdf(uploaded_file)
                elif file_type in ["docx", "doc"]:
                    st.session_state.content_df = self._parse_word(uploaded_file)
                else:
                    file_bytes = uploaded_file.getvalue()
                    csv_file = io.BytesIO(file_bytes)

                    success = False
                    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
                        if success:
                            break
                        for sep in [",", ";", "\t", "|"]:
                            try:
                                csv_file.seek(0)
                                df = pd.read_csv(csv_file, encoding=encoding, sep=sep, engine="python", on_bad_lines="skip")
                                if len(df.columns) > 1:
                                    st.session_state.content_df = df
                                    success = True
                                    break
                            except Exception:
                                continue

                    if not success:
                        st.error("无法解析CSV文件，请检查文件格式或尝试使用Excel格式(.xlsx)")
                        return

            df = st.session_state.content_df

            if df is None or len(df) == 0:
                callout("文件解析失败或无有效内容", type="error")
                return

            st.markdown("---")
            st.subheader("📋 选择脚本内容列")
            st.caption("系统需要知道哪一列包含脚本/文案内容")

            from utils.field_mapping import REVERSE_MAPPING
            auto_col = None
            for col in df.columns:
                col_lower = str(col).lower().strip()
                if col_lower in REVERSE_MAPPING and REVERSE_MAPPING[col_lower] == "脚本内容":
                    auto_col = col
                    break
                for keyword in ["脚本", "文案", "内容", "正文", "text", "content"]:
                    if keyword in col_lower:
                        auto_col = col
                        break
                if auto_col:
                    break

            if auto_col is None:
                for col in df.columns:
                    if df[col].dtype == "object" and df[col].notna().any():
                        sample = str(df[col].dropna().iloc[0]) if len(df[col].dropna()) > 0 else ""
                        if len(sample) > 20:
                            auto_col = col
                            break

            col_options = ["-- 请选择 --"] + list(df.columns)
            default_idx = col_options.index(auto_col) if auto_col and auto_col in col_options else 0

            selected_col = st.selectbox(
                "脚本内容列",
                col_options,
                index=default_idx,
                help="选择包含脚本/文案内容的列",
            )

            if selected_col != "-- 请选择 --":
                st.session_state.content_field_mapping = {"脚本内容": selected_col}
                st.success(f"✅ 已选择「{selected_col}」作为脚本内容列（共 {len(df)} 条记录）")

                with st.expander("预览数据"):
                    st.dataframe(df[[selected_col]].head(3), use_container_width=True)

                batch_btn = st.button(
                    f"开始批量分析（{len(df)} 条脚本）",
                    type="primary",
                    use_container_width=True,
                )

                if batch_btn:
                    self._handle_batch_analysis()
            else:
                st.warning("请选择包含脚本内容的列")
        else:
            st.session_state.content_df = None
            st.session_state.content_field_mapping = None

    def _handle_batch_analysis(self):
        from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
        from utils.field_mapping import normalize_columns

        if st.session_state.content_field_mapping is None:
            callout("请先完成字段映射", type="error")
            return

        mapping = st.session_state.content_field_mapping
        df = st.session_state.content_df

        if "脚本内容" not in mapping:
            callout("缺少必需的'脚本内容'字段映射", type="error")
            return

        df_normalized = normalize_columns(df, {k: v for k, v in mapping.items()})

        scripts = []
        seen_texts = set()
        for idx, row in df_normalized.iterrows():
            script_text = str(row.get("脚本内容", ""))
            if (not script_text
                    or not script_text.strip()
                    or script_text.lower() in ["nan", "none", ""]
                    or len(script_text.strip()) < 10):
                continue
            text_key = script_text.strip()[:50]
            if text_key in seen_texts:
                continue
            seen_texts.add(text_key)
            scripts.append({
                "script_text": script_text.strip(),
                "script_id": str(idx),
            })

        if not scripts:
            callout("未找到有效的脚本内容，请检查字段映射", type="error")
            return

        total = len(scripts)

        st.warning("⚠️ **分析中，请勿切换页面**，完成后将自动显示结果。")
        st.info(f"共 {total} 条脚本，并发处理中...")

        progress_bar = st.empty()
        status_text = st.empty()

        def _safe_progress(val, text):
            try:
                progress_bar.progress(val, text=text)
            except TypeError:
                progress_bar.progress(val)
                status_text.info(text)

        _safe_progress(0, f"准备分析 {total} 条脚本...")

        orchestrator = self._get_orchestrator()

        results = [None] * total
        completed = 0
        completed_lock = threading.Lock()

        max_workers = min(10, total)

        def analyze_one(index: int, script: dict):
            try:
                result = orchestrator.content_analyzer.analyze(
                    script.get("script_text", ""),
                    script_id=script.get("script_id")
                )
                return index, {"success": True, "data": result}
            except Exception as e:
                logger.error(f"内容分析失败 (item {index+1}/{total}): {e}")
                return index, {"success": False, "error": str(e)}

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(analyze_one, i, scripts[i]): i
                for i in range(total)
            }
            for future in as_completed(futures):
                try:
                    idx, result = future.result(timeout=300)
                except TimeoutError:
                    logger.error("内容分析超时 (item in batch)")
                    idx = futures[future]
                    result = {"success": False, "error": "分析超时(5分钟)，请检查网络或API状态"}
                except Exception as exc:
                    logger.error(f"内容分析线程异常: {exc}")
                    idx = futures[future]
                    result = {"success": False, "error": str(exc)}
                results[idx] = result
                with completed_lock:
                    completed += 1
                pct = int(completed / total * 100)
                _safe_progress(pct / 100, f"分析中 {completed}/{total} ({pct}%)")
                status_text.info(f"⏳ 已完成 {completed}/{total} 条")

        successful_results = [r["data"] for r in results if r and r.get("success")]
        if successful_results:
            orchestrator.db.save_content_analyses_batch(successful_results)
            _safe_progress(1.0, f"分析完成！共保存 {len(successful_results)} 条结果 ({int(completed / total * 100)}%)")
        else:
            _safe_progress(1.0, "分析完成，但无成功结果")

        state = {
            "results": results,
            "total": total,
        }

        self._show_batch_results(state)

    def _show_batch_results(self, state: dict):
        results = state.get("results", [])
        total = state.get("total", len(results))
        success_count = sum(1 for r in results if r.get("success"))
        fail_count = sum(1 for r in results if not r.get("success"))

        try:
            db_count = len(self._get_orchestrator().db.get_all_content_analyses())
        except Exception:
            db_count = -1

        msg = f"批量分析完成！成功 {success_count}/{total} 条"
        if fail_count > 0:
            msg += f"（{fail_count} 条失败）"
        if db_count >= 0:
            msg += f" | 数据库共 {db_count} 条记录"
        callout(msg, type="success", icon="✅")

        divider()
        st.subheader("✅ 当前完成结果")

        for i, r in enumerate(results):
            if r.get("success"):
                data = r.get("data", {})
                analysis = data.get("analysis", data)
                cid = data.get("content_id", "")[:8]
                score = analysis.get("content_score", "N/A")
                with st.expander(
                    f"#{i+1} [{cid}] - 评分 {score}/10"
                ):
                    self._display_analysis(analysis)
            else:
                with st.expander(f"脚本 #{i+1} - 分析失败"):
                    st.error(r.get("error", "未知错误"))

        st.session_state.content_df = None
        st.session_state.content_field_mapping = None
        if "batch_state" in st.session_state:
            del st.session_state.batch_state
        st.rerun()

    def _parse_pdf(self, uploaded_file) -> "pd.DataFrame":
        """解析 PDF 文件，提取文本内容

        默认按页提取文本，用户可选择表格模式
        """
        import pandas as pd

        # 让用户选择解析模式
        parse_mode = st.radio(
            "PDF 解析模式",
            options=["按页提取（推荐）", "表格提取"],
            horizontal=True,
            help="按页提取：每页作为一个脚本；表格提取：尝试识别PDF中的表格结构",
        )

        try:
            if parse_mode == "表格提取":
                return self._parse_pdf_as_table(uploaded_file)
            else:
                return self._parse_pdf_as_text(uploaded_file)
        except Exception as e:
            st.error(f"PDF 解析失败: {str(e)}")
            return pd.DataFrame()

    def _parse_pdf_as_text(self, uploaded_file) -> "pd.DataFrame":
        """按页提取 PDF 文本"""
        import pandas as pd

        try:
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                texts = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append({
                            "脚本内容": text.strip(),
                            "页码": i + 1,
                        })
        except ImportError:
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(uploaded_file)
            texts = []
            for i, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    texts.append({
                        "脚本内容": text.strip(),
                        "页码": i + 1,
                    })

        if not texts:
            st.warning("PDF 文件未提取到有效文本内容")
            return pd.DataFrame()

        df = pd.DataFrame(texts)
        st.success(f"✅ 成功从 PDF 提取 {len(df)} 页内容")
        return df

    def _parse_pdf_as_table(self, uploaded_file) -> "pd.DataFrame":
        """尝试从 PDF 提取表格"""
        import pandas as pd

        try:
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                all_tables = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        cleaned = []
                        for row in table:
                            cleaned_row = [(cell or "").strip() for cell in row]
                            if any(cell for cell in cleaned_row):
                                cleaned.append(cleaned_row)
                        if cleaned:
                            all_tables.extend(cleaned)
        except ImportError:
            st.error("表格模式需要 pdfplumber 库，当前不可用，请使用按页提取模式")
            return pd.DataFrame()

        if not all_tables or len(all_tables) <= 1:
            st.warning("未检测到有效表格，请尝试按页提取模式")
            return pd.DataFrame()

        # 处理不规则表格（列数不一致）
        max_cols = max(len(row) for row in all_tables)
        normalized = []
        for row in all_tables:
            if len(row) < max_cols:
                row = row + [""] * (max_cols - len(row))
            normalized.append(row[:max_cols])

        headers = normalized[0]
        df = pd.DataFrame(normalized[1:], columns=headers)
        df = df.dropna(how="all")
        df = df[~df.apply(lambda row: all(str(v).strip() in ["", "nan", "None"] for v in row), axis=1)]

        if len(df) == 0:
            st.warning("表格数据为空，请尝试按页提取模式")
            return pd.DataFrame()

        st.success(f"✅ 成功从 PDF 表格提取 {len(df)} 条记录")
        return df

    def _parse_word(self, uploaded_file) -> "pd.DataFrame":
        """解析 Word 文件，提取文本内容

        智能识别模式：
        1. 如果文档包含表格，提取表格数据（每行一条记录）
        2. 如果没有表格，按段落提取（合并连续短段落）
        """
        import pandas as pd
        from docx import Document

        try:
            doc = Document(uploaded_file)

            # 检查是否有表格
            has_tables = len(doc.tables) > 0

            if has_tables:
                parse_mode = st.radio(
                    "Word 解析模式",
                    options=["表格提取（推荐）", "段落提取"],
                    horizontal=True,
                    help="表格提取：提取文档中的表格数据；段落提取：按段落提取文本",
                )
            else:
                parse_mode = "段落提取"

            if parse_mode == "表格提取":
                return self._parse_word_as_table(doc)
            else:
                return self._parse_word_as_paragraphs(doc)

        except Exception as e:
            st.error(f"Word 解析失败: {str(e)}")
            return pd.DataFrame()

    def _parse_word_as_table(self, doc) -> "pd.DataFrame":
        """从 Word 文档提取表格数据"""
        import pandas as pd

        tables_data = []
        for table in doc.tables:
            for row in table.rows:
                # 去重合并单元格
                seen_texts = set()
                row_data = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in seen_texts:
                        seen_texts.add(text)
                        row_data.append(text)
                if any(row_data):
                    tables_data.append(row_data)

        if not tables_data or len(tables_data) <= 1:
            st.warning("未检测到有效表格，请尝试段落提取模式")
            return pd.DataFrame()

        # 处理不规则表格（列数不一致）
        max_cols = max(len(row) for row in tables_data)
        # 补齐短行
        normalized = []
        for row in tables_data:
            if len(row) < max_cols:
                row = row + [""] * (max_cols - len(row))
            normalized.append(row[:max_cols])  # 截断超长行

        headers = normalized[0]
        df = pd.DataFrame(normalized[1:], columns=headers)
        df = df.dropna(how="all")
        df = df[~df.apply(lambda row: all(str(v).strip() in ["", "nan", "None"] for v in row), axis=1)]

        if len(df) == 0:
            st.warning("表格数据为空，请尝试段落提取模式")
            return pd.DataFrame()

        st.success(f"✅ 成功从 Word 表格提取 {len(df)} 条记录")
        return df

    def _parse_word_as_paragraphs(self, doc) -> "pd.DataFrame":
        """从 Word 文档提取段落文本（智能合并连续短段落）"""
        import pandas as pd

        paragraphs = []
        buffer = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if buffer and len(buffer) > 20:
                    paragraphs.append(buffer.strip())
                buffer = ""
                continue

            if len(text) < 30 and len(buffer) < 200:
                buffer += text
            else:
                if buffer and len(buffer) > 20:
                    paragraphs.append(buffer.strip())
                buffer = text

        if buffer and len(buffer) > 20:
            paragraphs.append(buffer.strip())

        if not paragraphs:
            st.warning("Word 文档未提取到有效文本内容")
            return pd.DataFrame()

        df = pd.DataFrame({"脚本内容": paragraphs})
        st.success(f"✅ 成功从 Word 提取 {len(df)} 个段落")
        return df

    def _display_result(self, result: dict):
        """展示单个分析结果"""
        analysis = result["analysis"]
        content_id = result["content_id"]

        callout(f"分析完成！ID: {content_id[:8]}...", type="success", icon="✅")

        # 核心指标
        col1, col2, col3 = st.columns(3)
        with col1:
            metric_card(
                title="综合评分",
                value=f"{analysis.get('content_score', 'N/A')}/10",
                subtitle="综合转化潜力评分",
                icon="📊",
                border_color="#6366F1",
            )
        with col2:
            metric_card(
                title="Hook强度",
                value=f"{analysis.get('hook_strength', 'N/A')}/10",
                subtitle="开场钩子吸引力",
                icon="🎬",
                border_color="#10B981",
            )
        with col3:
            metric_card(
                title="CTA清晰度",
                value=f"{analysis.get('cta_clarity', 'N/A')}/10",
                subtitle="行动号召清晰程度",
                icon="📢",
                border_color="#F59E0B",
            )

        divider()
        self._display_analysis(analysis)

    def _display_analysis(self, analysis: dict):
        """展示分析详情"""
        render_content_analysis_details(analysis)

        divider()

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### 话题标签")
            tags = analysis.get("topic_tags", [])
            render_tags(tags)

        with col2:
            st.markdown("#### 目标受众")
            st.write(analysis.get("target_audience", "未知"))

        divider()

        render_list_section("核心卖点", analysis.get("key_selling_points", []))
        render_list_section("改进建议", analysis.get("improvement_suggestions", []))

    def _render_history(self):
        """展示历史分析记录"""
        st.subheader("历史分析记录")
        try:
            # 分页参数
            page_size = 10
            page = self._get_current_page("content_history")

            # 获取总数
            total_count = self._get_orchestrator().db.get_content_analyses_count()

            if total_count == 0:
                empty_state(
                    title="暂无历史记录",
                    description="去上方输入脚本开始分析吧！",
                    icon="📚",
                )
                return

            # 获取当前页数据
            offset = page * page_size
            records = self._get_orchestrator().db.get_all_content_analyses(
                limit=page_size, offset=offset
            )

            if not records:
                empty_state(
                    title="暂无历史记录",
                    description="没有找到匹配的记录。",
                    icon="📚",
                )
                return

            for record in records:
                analysis = record.get("analysis_json", {})
                score = analysis.get("content_score", "N/A")
                hook_type = analysis.get("hook_type", "未知")
                raw_text = record.get("raw_text", "")[:60] + "..."
                rid = record.get("id", "")[:8]

                with st.expander(
                    f"[{rid}] 评分 {score}/10 | {hook_type} | {record['created_at'][:10]}"
                ):
                    st.caption(raw_text)
                    divider()
                    self._display_analysis(analysis)

            # 分页控制
            self._render_pagination(total_count, page_size, "content_history")

        except Exception as e:
            callout(f"加载历史记录失败: {str(e)}", type="error")


def render_content_analysis():
    """页面入口函数"""
    page = ContentAnalysisPage()
    page.render()
